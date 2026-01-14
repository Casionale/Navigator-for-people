import json
import os
import random
from datetime import datetime

import pandas
import pandas as pd
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from pathlib import Path


class bcolors:
    HEADER = '\033[95m'
    OKBLUE: str = '\033[94m'
    OKCYAN: str = '\033[96m'
    OKGREEN: str = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

class rgbcolors:
    @staticmethod
    def color(r, g, b):
        return '\033[38;2;{0};{1};{2}m '.format(r,g,b)
    @staticmethod
    def end():
        return '\033[0m'

class progressBar():
    size = 30
    filled = '█'
    unfilled = '-'

    def __init__(self):
        self.size = 30
        self.filled = '█'
        self.unfilled = '-'

    def getPB(self, all, progress):
        percent = int((progress * 100) / all)
        filled_count = int((self.size / 100) * percent)
        fil = str(self.filled*filled_count) + str(self.unfilled * (self.size-filled_count))
        return "{0}% {1} 100%".format(percent, fil)

url = "https://booking.dop29.ru/api/user/login"
user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0"

dir = os.getcwd()

file_login = open(dir + '\\login.ini', 'r')
str_login = file_login.read().split('\n')
email = str_login[0]
password = str_login[1]
YEAR = str_login[2]
OUTPUT_DIR = Path(dir) / 'output'

if not OUTPUT_DIR.exists():
    os.makedirs(OUTPUT_DIR)

def get_save_path(filename):
    global OUTPUT_DIR
    now = datetime.now()
    str_now = now.strftime('%m-%d-%y')

    folder_path = OUTPUT_DIR / str_now

    if not folder_path.exists():
        os.makedirs(folder_path)

    return folder_path / filename

session = requests.Session()
r = session.post(url, headers={
    'Host': 'booking.dop29.ru',
    'User-Agent': user_agent,
    'Accept': '*/*',
    'Accept-Language': 'ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3',
    'Accept-Encoding': 'gzip, deflate, br',
    'Content-Type': 'application/json',
    'X-Requested-With': 'XMLHttpRequest',
    'Content-Length': '63',
    'Origin': 'https://booking.dop29.ru',
    'DNT': '1',
    'Connection': 'keep-alive',
    'Referer': 'https://booking.dop29.ru/admin/',
    'Sec-Fetch-Dest': 'empty',
    'Sec-Fetch-Mode': 'cors',
    'Sec-Fetch-Site': 'same-origin',
    'TE': 'trailers',
}, data='{"email": "' + email + '", "password": "' + password + '"}')



session.headers.update({'Referer': 'https://booking.dop29.ru/admin/'})
session.headers.update({'User-Agent': user_agent})

text_buf = r.text
json_string = json.loads(text_buf)

if json_string['err_code'] == 400:
    print(json_string['errors'][0]['msg'])
    exit(1)
else:
    print('Авторизация удалась походу')

access_token = json_string['data']['access_token']
expired_at = json_string['data']['expired_at']
refresh_token = json_string['data']['refresh_token']

user = json_string['data']['user']

MAX_GROUPS_COUNT = 500

headers = {
    'Host': 'booking.dop29.ru',
    'User-Agent': user_agent,
    'Accept': '*/*',
    'Accept-Language': 'ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3',
    'Accept-Encoding': 'gzip, deflate, br',
    'Authorization': 'Bearer ' + access_token,
    'X-REQUEST-ID': '7bd411c3-54ce-4bba-9ee1-7c5091da6d1a',
    'X-Requested-With': 'XMLHttpRequest',
    'DNT': '1',
    'Connection': 'keep-alive',
    'Referer': 'https://booking.dop29.ru/admin/',
    'Cookie': 'io=lVluIaMvSTa4ImFmB5C9',
    'Sec-Fetch-Dest': 'empty',
    'Sec-Fetch-Mode': 'cors',
    'Sec-Fetch-Site': 'same-origin',
    'TE': 'trailers'
}

new_url = 'https://booking.dop29.ru/api/rest/eventGroups?_dc=1641896017213&page=1&start=0&length=25&extFilters=[{"property":"is_deleted","value":"0","comparison":"eq"},{"property":"event.is_deleted","value":"N","comparison":"eq"}]&format=attendance&length='+str(MAX_GROUPS_COUNT)
r = session.get(new_url, headers=headers)

b = json.loads(r.text)
groups = b['data']

if int(b['recordsFiltered']) > len(groups):
    print("Загружено {0} из {1}".format(len(groups), int(b['recordsFiltered'])))

    new_url = 'https://booking.dop29.ru/api/rest/eventGroups?_dc=1641896017213&page=1&start=0&length=25&extFilters=[{"property":"is_deleted","value":"0","comparison":"eq"},{"property":"event.is_deleted","value":"N","comparison":"eq"}]&format=attendance&length=' + str(
        MAX_GROUPS_COUNT)+'&page=2&start='+str(len(groups))
    r = session.get(new_url, headers=headers)

    b = json.loads(r.text)
    groups.extend(b['data'])

    print("Загружено {0} из {1}".format(len(groups), int(b['recordsFiltered'])))

i = -1


def printChildren():
    global new_url, r, b, i
    print('Выбрана группа ' + groups[g_inp]['program_name'] + ' ' + groups[g_inp]['name'])
    year = YEAR
    list_childrens = get_childrens()
    f = open(get_save_path(f"Список группы {groups[g_inp]['program_name']} {groups[g_inp]['name']}.txt"), 'w', encoding="utf-8")
    for c in list_childrens:
        line = (f'{c['kid_last_name']} {c['kid_first_name']} {c['kid_patro_name']}\t'
                f'{c['kid_birthday'].replace('-', '.')}\t{c['kid_age']}\n')
        f.write(line)
    f.close()
    return 'Список группы ' + groups[g_inp]['program_name'] + ' ' + groups[g_inp]['name'] + ".txt"

def stat_of_ages(unique = False, confirmed = False, by_program_name = False, negative_groups = [], by_groups = False):
    global new_url, r, b, i, group_id_val
    ages = {0:0, 1:0, 2:0,3:0,4:0,5:0,6:0,7:0,8:0,9:0,10:0,11:0,12:0,13:0,14:0,15:0,16:0,17:0,18:0,19:0,20:0,21:0}
    stupid_girls_by_ages = {0:0, 1:0, 2:0,3:0,4:0,5:0,6:0,7:0,8:0,9:0,10:0,11:0,12:0,13:0,14:0,15:0,16:0,17:0,18:0,19:0,20:0,21:0}
    sum_girls = 0
    sum_childs = 0
    error_sex = 0
    error_child = []

    unique_childs = set()
    repeated_childs = 0
    repeteds = 0
    repeated_childs_id = []

    ages_of_sections ={} #секция: [мальчики, девочки]
    for i in range(0, len(groups)):

        if len(negative_groups) > 0:
            negative_check = [ind for ind in negative_groups if ind in groups[i]['name'].lower()]
            if len(negative_check) != 0:
                print('\r Проигнорированна группа '+groups[i]['name'] + '\n')
                continue


        g_inp = i
        group_id_val = groups[i]['id']
        event_id = groups[i]['event_id']

        #print('Выбрана группа ' + groups[g_inp]['program_name'] + ' ' + groups[g_inp]['name'])
        pb = progressBar()
        print("\r"+pb.getPB(all=len(groups), progress=i)+' Выбрана группа ' + groups[g_inp]['program_name'] + ' ' + groups[g_inp]['name'], end=' ')
        list_childrens = get_childrens()
        if by_program_name:
            section = groups[i]['program_name']
        else:
            if by_groups:
                section = groups[i]['program_name'] + " " + groups[i]['name']
            else:
                section = get_section(event_id)

        if section not in ages_of_sections:
            ages_of_sections[section] = ([0],[0],{0:0, 1:0, 2:0,3:0,4:0,5:0,6:0,7:0,8:0,9:0,10:0,11:0,12:0,13:0,14:0,15:0,16:0,17:0,18:0,19:0,20:0,21:0,22:0,23:0},
                                         {0:0, 1:0, 2:0,3:0,4:0,5:0,6:0,7:0,8:0,9:0,10:0,11:0,12:0,13:0,14:0,15:0,16:0,17:0,18:0,19:0,20:0,21:0,22:0,23:0}) #Мальчики, девочки, всего по годам, из них дев по годам

        iterator_childrens = 0


        for c in list_childrens:
            if unique:
                if c['kid_id'] in unique_childs:
                    repeteds += 1
                    if c['kid_id'] not in repeated_childs_id:
                        repeated_childs += 1
                        repeated_childs_id.append(c['kid_id'])
                    continue
                else:
                    unique_childs.add(c['kid_id'])

            ages[c['kid_age']] += 1
            ages_of_sections[section][2][c['kid_age']] += 1
            c_info = get_all_info_child(c['kid_id'])
            sum_childs += 1

            try:
                if c_info['sex'] == 'W':
                    sum_girls += 1
                    ages_of_sections[section][1][0] += 1
                    stupid_girls_by_ages[c['kid_age']]+=1
                    ages_of_sections[section][3][c['kid_age']] += 1
                else:
                    ages_of_sections[section][0][0] += 1
            except:
                error_sex += 1
                error_child.append("{0} {1} {2} {3}".format(c['kid_last_name'],c['kid_first_name'],c['kid_patro_name'],c['kid_birthday']))

            iterator_childrens += 1
            print("\r" + pb.getPB(all=len(groups), progress=i) + ' Выбрана группа ' + groups[g_inp][
                'program_name'] + ' ' + groups[g_inp]['name'], end=' ')
            print(pb.getPB(all=len(list_childrens), progress=iterator_childrens), end="")

    if not unique:
        f = open(get_save_path("Статистика по возрастам.txt"), "w", encoding='utf-8')
    else:
        f = open(get_save_path("Статистика по возрастам УНИКАЛЬНЫЕ.txt"), "w", encoding='utf-8')
    for i in range(0, 19):
        if ages[i] == 0:
            continue
        else:
            f.write(str(i) + " лет " + str(ages[i]) + f" человек; {stupid_girls_by_ages[i]} из них девочек \n")

    f.write("Всего: {0}, из них девочек: {1}".format(sum_childs, sum_girls))
    f.write("\nНе удалось получить информацию у {0} человек:".format(error_sex))
    for l in error_child:
        f.write(l)
    for key, value in ages_of_sections.items():
        f.write("\n\nНаправленность: {0}, М {1}, Ж {2}\n\n".format(key, value[0][0], value[1][0]))

        for i in range(0, 19):
            if value[2][i] == 0:
                continue
            else:
                f.write(str(i) + " лет " + str(value[2][i]) + " человек, из них девочек  "+str(value[3][i])+"\n")

    if unique:
        f.write("\nПовторов всего: " + str(repeteds))
        f.write("\nПовторов детей: " + str(repeated_childs))
        f.write('\n'+str(repeated_childs_id))

    f.close()

def get_section(event_id):
    new_url = 'https://booking.dop29.ru/api/rest/events/{0}?_dc=1705994318220'.format(event_id)
    r = session.get(new_url, headers=headers)
    b = json.loads(r.text)
    try:
        return b['data'][0]['section']
    except:
        return []

def get_childrens():
    new_url = 'https://booking.dop29.ru/api/attendance/members/get?_dc=1641896197594&page=1&start=0&length=25&extFilters=[{"property":"group_id","value":"' + str(
        group_id_val) + '"},{"property":"academic_year_id","value":"' + str(
        YEAR) + '"},{"property":"dateStart","value":"' + YEAR + '-12-01 00:00:00"},{"property":"dateEnd","value":"' + YEAR + '-12-31 23:59:59"}]'
    r = session.get(new_url, headers=headers)
    b = json.loads(r.text)
    list_childrens = b['data']
    new_list_childrens = []
    for i in range(0, len(list_childrens)):
        if list_childrens[i]['type_active'] == 1:
            new_list_childrens.append(list_childrens[i])
    return new_list_childrens

def get_childrens_by_group_id(group_id):
    new_url = 'https://booking.dop29.ru/api/attendance/members/get?_dc=1641896197594&page=1&start=0&length=25&extFilters=[{"property":"group_id","value":"' + str(
        group_id) + '"},{"property":"academic_year_id","value":"' + str(
        YEAR) + '"},{"property":"dateStart","value":"' + YEAR + '-12-01 00:00:00"},{"property":"dateEnd","value":"' + YEAR + '-12-31 23:59:59"}]'
    r = session.get(new_url, headers=headers)
    b = json.loads(r.text)
    list_childrens = b['data']
    new_list_childrens = []
    for i in range(0, len(list_childrens)):
        if list_childrens[i]['type_active'] == 1:
            new_list_childrens.append(list_childrens[i])
    return new_list_childrens

def get_all_info_child(id):
    new_url = "https://booking.dop29.ru/api/rest/kid/{0}?_dc=1704971612231".format(id)
    r = session.get(new_url, headers=headers)
    b = json.loads(r.text)
    try:
        child = b['data'][0]
        return child
    except:
        return []

def getFileListChildrensFromOrder(group):
    global g_inp, group_id_val
    g_inp = group

    group_id_val = group['id']

    new_url = 'https://booking.dop29.ru/api/rest/order?_dc=1695285515100&page=1&start=0&length=25&extFilters=[{"property":"fact_academic_year_id","value":'+YEAR+',"comparison":"eq"},{"property":"event_id","value":'+ group['event_id'] +',"comparison":"eq"},{"property":"fact_group_id","value":"' + str(group_id_val) + '","comparison":"eq"},{"property":"state","value":["approve"],"comparison":"in"}]'

    r = session.get(new_url, headers=headers)
    b = json.loads(r.text)
    list_childrens = b['data']
    list_names = []
    for children in list_childrens:
        url_child = 'https://booking.dop29.ru/api/rest/kid/'+children['kid_id']
        r = session.get(url_child, headers=headers)
        child = json.loads(r.text)['data'][0]
        list_names.append(child['last_name'] + " " + child['first_name'] + " " + child['patro_name'])

    file = open(get_save_path(f"Подтверждённые заявки {group['program_name']} {group['name']}.txt"),
                'w', encoding="utf-8")
    file.write("\n".join(list_names))
    file.close()

def getListChildrensFromOrder(group):

    group_id_val = group['id']

    new_url = 'https://booking.dop29.ru/api/rest/order?_dc=1695285515100&page=1&start=0&length=25&extFilters=[{"property":"fact_academic_year_id","value":'+YEAR+',"comparison":"eq"},{"property":"event_id","value":'+ group['event_id'] +',"comparison":"eq"},{"property":"fact_group_id","value":"' + str(group_id_val) + '","comparison":"eq"},{"property":"state","value":["approve"],"comparison":"in"}]'

    r = session.get(new_url, headers=headers)
    b = json.loads(r.text)
    list_childrens = b['data']
    return list_childrens

def getListChildrensFromOrderAnyGroups(groups_list):
    groups_list = groups_list.split(" ")
    for group in groups_list:
            for g in groups:
                if g == groups[int(group)]:
                    getFileListChildrensFromOrder(g)



#Типы занятий
#Практическая работа 30011
#Учебное 9419
#Дистанционное 3022
#
def close_day(date, theme, type, description):
    percentagevisits = 80
    global g_inp, group_id_val
    group_id_val = groups[int(g_inp)]['id']
    group_id = groups[int(g_inp)]['id']

    childrens = get_childrens()

    for c in childrens:
        visit = random.randint(0, 100)
        if visit < percentagevisits:
            visit = True
        else:
            visit = False
        kid_id = c['kid_id']
        day_url_POST = "https://booking.dop29.ru/api/attendance/save" #date group_id kid_id value:true
        payload = {'date':date,'group_id':group_id, 'kid_id':kid_id, 'value':visit}
        r = session.post(url=day_url_POST, headers=headers, data=json.dumps(payload))
        b = json.loads(r.text)
        #print(b["success"])

    KTP_url_POST = "https://booking.dop29.ru/api/event-group-lessons/upsert"# date group_id types:list description

    payload = json.loads(json.dumps({'data':{'date':date, 'group_id': group_id,'theme': theme ,'types': [str(int(type))], 'description': description}}))
    #payload = '{"data":{"group_id":"25849","date":"2023-09-05 00:00:00","types":["9732"]}}'
    r = session.post(url=KTP_url_POST, headers=headers, json=payload)
    b = json.loads(r.text)
    #print(b["success"])
    pass

def number_6(target_sum):
    global group_id_val
    problem_groups = []
    for g in groups:
        id = g['id']
        group_id_val = id
        childrens = get_childrens()
        if len(childrens) < target_sum and len(childrens) != 0:
            problem_groups.append("Группа: {0}, {1} человек!".format (g['program_name'] + " " + g['name'], len(childrens)))

    f = open(get_save_path("ПРОБЛЕМНЫЕ ГРУППЫ.txt"), "w")
    for g in problem_groups:
        f.write(g + '\n')
    f.close()

def find_duplicates():
    global group_id_val
    childs_and_groups = {}
    for g in groups:
        id = g['id']
        group_id_val = id
        childs = get_childrens()

        for child in childs:
            full_name = child['kid_last_name'] + ' ' + child['kid_first_name'] + ' ' + child['kid_patro_name']
            group_name = g['program_name'] + ' ' + g['name']

            if full_name in childs_and_groups:
                childs_and_groups[full_name].append(group_name)
            else:
                childs_and_groups[full_name] = []
                childs_and_groups[full_name].append(group_name)

            pass

    duplicated = {}
    for key in childs_and_groups:
        if len(childs_and_groups[key]) > 1:
            duplicated[key] = childs_and_groups[key]

    f = open(get_save_path("Дубликаты.txt"), "w")
    for key in duplicated:

        str_groups = '\n'
        for i in range(len(duplicated[key])):
            str_groups += '\t' + duplicated[key][i] + '\n'
        str_groups += '\n'

        f.write('Групп ' + str(len(duplicated[key])) + ' ' + key + ' ' + str_groups + '\n')

    f.write('\n\n Всего детей: ' + str(len(duplicated)))
    f.close()

def forced_child_adding(in_group = True):
    #Файл с детьми которых нужно зачислить в группы
    #Текстовый
    #ФИО таб часть program_name

    #Файл с детьми которых нужно зачислить в мероприятия
    #Текстовый
    #ИД мероприяти таб дата-время
    #ФИО таб описание

    all_childrens = []

    if in_group:
        for g in groups:
            childs = get_childrens_by_group_id(g['id'])
            all_childrens.extend(childs)
            also_childs = getListChildrensFromOrder(g)
            all_childrens.extend(also_childs)

    filename = input('Файл с детьми для добавления')
    f = open(filename, 'r', encoding='utf-8')
    rows = f.readlines()
    f.close()

    activity_id = 0
    date = ""
    if not in_group:
        parts = rows[0].split('\t')
        activity_id = int(parts[0])
        date = parts[1]
        del(rows[0])

    prerared_info = []
    for r in rows:
        parts = r.split('\t')
        if parts[0] != '':
            prerared_info.append(parts)

    actual_kids_for_adding = []
    for info in prerared_info:
        full_name = info[0]
        try:
            if info[1] is not None:
                adding_info = info[1]
            else:
                adding_info = ''
        except IndexError:
            adding_info = ''
        print('Поиск {0}'.format(full_name))
        target_url = f'https://booking.dop29.ru/api/rest/safe/kid?_dc=1714046462894&special=1&page=1&start=0&length=20&extFilters=[{{"property":"fio","value":"{full_name}","comparison":"manual","type":null}}]'
        r = session.get(url=target_url, headers=headers)
        b = json.loads(r.text)

        if b['err_code'] != 0 or len(b['data']) == 0:
            print('Не найдено!')
            continue
        os.system('cls')
        print('Найдены следующие дети:')

        for i in range(len(b['data'])):
            print(f"{i} {b['data'][i]['fio']} {b['data'][i]['birthday']} {b['data'][i]['approve_org_caption']}")

        if not in_group:
            print(f'Описание: {adding_info}')

        choose = int(input('Выберите индекс ребёнка для добавления или напишите -1 для пропуска'))
        if choose == -1:
            continue

        actual_kids_for_adding.append(b['data'][choose])

        if in_group:
            print("Веберите идекс группы для зачисления")
            for i in range(len(groups)):
                if adding_info.lower().rstrip() in groups[i]['program_name'].lower().rstrip():
                    print(f"{i} {groups[i]['program_name']} {groups[i]['id']} {groups[i]['name']}")

            if b['data'][choose]['id'] in [kid['kid_id'] for kid in all_childrens]:
                print('УЖЕ ДОБАВЛЕН В КАКОЙ-ТО ГРУППЕ')

            group_index = int(input('Выбранная группа: '))
            if group_index != -1:
                print(f"{groups[group_index]['program_name']} {groups[group_index]['id']} {groups[group_index]['name']}")
                adding_order(b['data'][choose], groups[group_index])
            else:
                print('Пропущен!')
        else:
            adding_activity_order(b['data'][choose], activity_id, date)
        pass


def adding_order(child, group):
    json_string = {"data":
                {"event_id":group['event_id'],
                 "state":"initial",
                 "certificate_number":"нет",
                 "decree_enrollment_number":"нет",
                 "decree_deduction_number":"нет",
                 "program_is_pfdod":False,
                 "kid_is_approved":False,
                 "is_online_payments_allowed":False,
                 "academic_year_id":YEAR,
                 "certificate_certificate_number":"",
                 "rpgu_deadline_date":None,
                 "kid_birthday":None,
                 "deadline":None,
                 "created_ts":None,
                 "date_enroll":None,
                 "date_deduct":None,
                 "rpgu_overdue_deadline":False,
                 "group_id":f"{group['id']}",
                 "kid_id":f"{child['id']}",
                 "site_user_id":f"{child['site_user_id']}"
                 }}
    new_url = "https://booking.dop29.ru/api/rest/order?_dc=1714059370156"
    payload = json.loads(json.dumps(json_string))
    r = session.post(url=new_url, headers=headers, json=payload)
    b = json.loads(r.text)
    if b['err_code'] == 0:
        print('УСПЕХ!')
    else:
        b['errors'][0]['msg']

def adding_activity_order(child, activity_id, date, state = 'approve'):
    json_string = {"data":
                       {
                           "activity_id": activity_id,
                           "date": date,
                           "site_user_id": f"{child['site_user_id']}",
                           "kid_id": f"{child['id']}",
                           "state": state
                       }}
    new_url = "https://booking.dop29.ru/api/rest/activityOrder?_dc=1714066258891"
    payload = json.loads(json.dumps(json_string))
    r = session.post(url=new_url, headers=headers, json=payload)
    b = json.loads(r.text)
    if b['err_code'] == 0:
        print('УСПЕХ!')
    else:
        b['errors'][0]['msg']

def to_study_from_approve():
    global groups

    target_url = f'https://booking.dop29.ru/api/rest/order?_dc=1714828457975&page=1&start=0&length={MAX_GROUPS_COUNT}&extFilters=[{{"property":"fact_academic_year_id","value":{YEAR},"comparison":"eq"}}]'
    r = session.get(url=target_url, headers=headers)
    b = json.loads(r.text)
    approving = [a for a in b['data'] if a['state_grid'] == 'approve']
    print(f'Найдено {len(approving)} подтверждённых заявок для подтверждения обучения')
    date_signing = "" #Дата приказа
    date_start = "" #Начало обучения
    decree_number = "" #Номер приказа
    financing_source = "1" #Бюджет 1
    id = "" #id заявки ребёнка

    print('='*20)
    date_signing = input('Дата приказа в формате ГГ-ММ-ДД! ')
    date_start = input('Дата начала обучения в формате ГГ-ММ-ДД!')
    decree_number = input('Номер приказа ')

    for a in approving:
        g = next((x for x in groups if x['id'] == a['group_id']),  None)
        if g is None:
            continue

        print(f"Ребёнок {a['kid_last_name']} {a['kid_first_name']} {g['program_name']} {g['name']}")
        choose = input('1 - Принять; 0 - пропустить;\n')
        if choose == '1':
            id = a['id']

            target_url = f'https://booking.dop29.ru/api/studyRequest'
            json_string = {"data":
                {
                    "comment": "",
                    "date_signing": date_signing,
                    "date_start": date_start,
                    "decree_number": decree_number,
                    "financing_source": "1",
                    "id":id
                }}
            payload = json.loads(json.dumps(json_string))
            r = session.post(url=target_url, headers=headers, json=payload)
            b = json.loads(r.text)
            if b['err_code'] == 0:
                print('УСПЕХ!')
            else:
                b['errors'][0]['msg']

        else:
            continue

    pass

def generateDiagnostic(group, existing = True):
    global group_id_val, groups
    global diagnostics_sums

    group_id_val = groups[int(group)]['id']
    childrens = get_childrens()
    # Высокий уровень 30%, средний 70% для выходной
    list_fio = [f"{c['kid_last_name']} {c['kid_first_name']} {c['kid_patro_name']}" for c in childrens]

    table = []
    summary = generate_data(diagnostics_sums, group, groups, list_fio, table, existing=existing)

    header_table = [['#п/п', 'ФИО обучающегося', 'уровень знаний', '', ''],
                    ['',     '',                 'низкий',         'средний', 'высокий']]

    header_table.extend(table)
    table = header_table

    doc = create_document()

    if existing:
        parts = [(f"Выходная диагностика {groups[int(group)]['program_name']} {groups[int(group)]['name']}", True)]
    else:
        parts = [(f"Входная диагностика {groups[int(group)]['program_name']} {groups[int(group)]['name']}", True)]
    add_paragraph(doc, parts, font_size=14, alignment='center')

    t = add_table(doc, table)
    merge_cells_with_content(t, 0, 0, 1, 0)
    merge_cells_with_content(t, 0, 1, 1, 1)
    merge_cells_with_content(t, 0, 2, 0, 4)

    parts = [(f"Итого: {summary[0]} низкий, {summary[1]} средний, {summary[2]} высокий", False)]
    add_paragraph(doc, parts, font_size=14, alignment='center')

    add_page_break(doc)

    parts = [(f"Критерии оценки", True)]
    add_paragraph(doc, parts, font_size=14, alignment='center')

    table = [['Уровень знаний и умений', 'Низкий уровень', 'Средний уровень', 'Высокий уровень'],
             ['Теоретические знания', 'Определяются по результатам собеседования', '', ''],
             ['Практические умения и навыки', 'Ребенок не смог выполнить задание без помощи педагога или работал самостоятельно, но задание выполнено не верно', 'Задание выполнено хорошо, но ребенок задавал вопросы в процессе выполнения', 'Задание выполнено самостоятельно, быстро и качественно'],
             ['Личностные качества', 'Определяются в результате педагогического наблюдения в процессе выполнения задания', '', ''],
             ]

    t = add_table(doc, table)
    merge_cells_with_content(t, 1, 1, 1, 3)
    merge_cells_with_content(t, 3, 1, 3, 3)

    parts = [(f"Форма определения уровня освоения программы:", True), (" педагогическое наблюдение, собеседование, анализ практической работы, результат проекта.", False)]
    add_paragraph(doc, parts, font_size=14, alignment='justify')

    if existing:
        save_document(doc,
                  get_save_path(f"Выходная диагностика {groups[int(group)]['program_name']} {groups[int(group)]['name']}.docx"))
    else:
        save_document(doc,
                      get_save_path(f"Входная диагностика {groups[int(group)]['program_name']} {groups[int(group)]['name']}.docx"))


def generate_data(diagnostics_sums, group, groups, list_fio, table, existing = True):
    summary = [0, 0, 0]
    for i in range(len(list_fio)):
        r = random.randint(0, 100)
        if r >= 70:
            if existing:
                table.append([i + 1, list_fio[i], "", "", "+"])
            else:
                table.append([i + 1, list_fio[i], "", "+", ""])
            if f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}" not in \
                    diagnostics_sums.keys():
                diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"] = \
                    {'high': 0, 'middle': 0}
            diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"]['high'] += 1

            summary[2 if existing else 1] += 1
        else:
            if existing:
                table.append([i + 1, list_fio[i], "", "+", ""])
            else:
                table.append([i + 1, list_fio[i], "+", "", ""])

            if f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}" not in \
                    diagnostics_sums.keys():
                diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"] = \
                    {'high': 0, 'middle': 0}
            diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"]['middle'] += 1

            summary[1 if existing else 0] += 1
    return summary


def getDiagnostics(groups, existing = True):
    global diagnostics_sums
    diagnostics_sums = {}
    if ' ' in groups:
        groups = groups.split(' ')
        for group in groups:
            generateDiagnostic(int(group)-1, existing=existing)
    else:
        generateDiagnostic(int(groups)-1, existing=existing)
    f = open(get_save_path('Диагностика суммы.txt'), 'w', encoding='utf-8')
    for key, value in diagnostics_sums.items():
        f.write(f"{key} Высокий: {value['high']} Средний: {value['middle']}\n")
    f.close()
    pass


def create_document():
    """Создаёт новый документ Word."""
    return Document()


def add_paragraph(doc, parts, font_size=14, alignment=None):
    """
    Добавляет абзац с текстом в документ, поддерживая отдельные жирные слова.
    :param doc: документ
    :param parts: список кортежей (текст, жирный), например, [("Hello", False), ("World", True)]
    :param font_size: размер шрифта
    :param alignment: выравнивание текста (None, 'center', 'left', 'right', 'justify')
    """
    paragraph = doc.add_paragraph()
    for text, bold in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)

    if alignment:
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            # Выравнивание по обеим сторонам (приближение к выравниванию по ширине)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return paragraph


def set_cell_border(cell, **kwargs):
    """
    Устанавливает границы для ячейки.
    :param cell: ячейка таблицы
    :param kwargs: параметры границ
    """
    tcPr = cell._element.get_or_add_tcPr()

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = tcPr.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tcPr.append(border)
        for attr, value in kwargs.items():
            border.set(qn(f'w:{attr}'), str(value))


def add_table(doc, data):
    """
    Добавляет таблицу в документ с границами.
    :param doc: документ
    :param data: список списков, представляющий строки и столбцы таблицы
    """
    table = doc.add_table(rows=len(data), cols=len(data[0]))

    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table_cell = table.cell(i, j)
            table_cell.text = str(cell)
            set_cell_border(table_cell, val="single", sz="4", space="0", color="000000")

    # Автоматически подгоняем ширину столбцов под текст
    for col in table.columns:
        max_length = max(len(cell.text) for cell in col.cells)
        for cell in col.cells:
            cell.width = Inches(0.15 * max_length)  # Установка ширины в зависимости от длины текста
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def merge_cells_with_content(table, start_row, start_col, end_row, end_col):
    """
    Объединяет ячейки таблицы от start_row, start_col до end_row, end_col, если одна из ячеек заполнена, а другие пустые.
    :param table: таблица
    :param start_row: начальная строка
    :param start_col: начальный столбец
    :param end_row: конечная строка
    :param end_col: конечный столбец
    """
    # Проверка содержимого ячеек
    content = None
    for row in range(start_row, end_row + 1):
        for col in range(start_col, end_col + 1):
            cell_text = table.cell(row, col).text.strip()
            if cell_text:
                if content and content != cell_text:
                    raise ValueError("Конфликтующее содержимое ячеек")
                content = cell_text

    # Объединение ячеек
    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)
    start_cell.merge(end_cell)

    if content:
        start_cell.text = content


def center_text(cell):
    """
    Выравнивает текст в ячейке по центру.
    :param cell: ячейка
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(14)  # Установка размера шрифта


def add_line_break(doc):
    """Добавляет разрыв строки в документ."""
    doc.add_paragraph().add_run().add_break()


def add_page_break(doc):
    """Добавляет разрыв страницы в документ."""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def save_document(doc, filename):
    """Сохраняет документ под указанным именем файла."""
    doc.save(filename)

def child_search_online():
    print('Для выхода введите #')
    while True:
        full_name = input("Введите ФИО: ")

        if full_name == '#':
            break

        b = search_child_by_fio(full_name)

        os.system('cls')
        print('Найдены следующие дети:')

        for i in range(len(b['data'])):
            print(f"{i} {b['data'][i]['fio']} {b['data'][i]['birthday']} {b['data'][i]['approve_org_caption']}")


def search_child_by_fio(full_name):
    target_url = f'https://booking.dop29.ru/api/rest/safe/kid?_dc=1714046462894&special=1&page=1&start=0&length=20&extFilters=[{{"property":"fio","value":"{full_name}","comparison":"manual","type":null}}]'
    r = session.get(url=target_url, headers=headers)
    b = json.loads(r.text)
    if b['err_code'] != 0 or len(b['data']) == 0:
        print('Не найдено!')
    return b

FILTER = False
diagnostics_sums = {}

filter_choise = int(input("Режим фильтра 0 - нет, 1 - да: "))

if filter_choise == 1:
    f = open("groups.ini")
    id_filters = f.read().splitlines()
    f.close()

    filtred_groups = []

    for g in groups:
        pass
        if g['id'] in id_filters:
            filtred_groups.append(g)
    groups = filtred_groups

    filter_teachers = int(input("Фильтр по педагогам 0 - нет, 1 - да: "))

    if filter_teachers == 1:
        teachers = {}
        for g in groups:
            if teachers.get(g["teacher"]) is not None:
                teachers[g["teacher"]].append(g)
            else:
                teachers[g["teacher"]] = []
                teachers[g["teacher"]].append(g)

        keys = []
        for t in teachers:
            keys.append(t)

        print("Список педагогов")
        for i in range(0, len(keys)):
            print("{0} {1}".format(i, keys[i]) )

        teacher_groups = input("Группы каких педагогов выбрать? Можно через пробел указать ").split(' ')

        groups = []

        for t in teacher_groups:
            groups.extend(teachers[keys[int(t)]])

        print("Выбраны {0} групп".format(len(groups)))


while True:
    os.system("")
    choose = input(bcolors.OKGREEN + 'МЕНЮ' + bcolors.ENDC +'\n'
                   '0 Печать информации детей\n'
                   '3 Печать статистики по возрастам\n'
                   '4 Печать списка из заявок (Когда зачисления ещё нет, но хочется получить список)\n'
                   '5 ! Внести в навигатор свои{0}грязные{1} буквы\n'.format(rgbcolors.color(255, 128, 0),
                                                                             rgbcolors.end()) +
                   '6 Найти проблемные группы\n'
                   '7 Найти дубликаты детей\n'
                   '{0}8 По возрастам и уникальные{1}\n'.format(rgbcolors.color(127, 255, 212),
                                                                rgbcolors.end()) +
                   '9 количество детей по программам\n'
                   '{0}10 принудительная заявка детей в группу{1}\n'.format(rgbcolors.color(198, 144, 53),
                                                                            rgbcolors.end()) +
                   '{0}11 принудительное зачисление детей в мероприятие{1}\n'.format(rgbcolors.color(198, 144, 53),
                                                                                     rgbcolors.end()) +
                   '12 Принять на обучение\n'
                   '13 Генерировать выходную диагностику\n'
                   '14 Поиск детей онлайн по ФИО\n'
                   '15 Генерировать входную диагностику\n'
                   '# Вернуться в главное меню (во всей программе)')

    i = 0

    if choose == '0':
        print('Группы')
        for g in groups:
            i = i + 1
            print(str(i) + ' ' + g['program_name'] + ' ' + g['id'] + " " + g['name'])
        print('Какую группу вывести на печать? ')

        input_str = input()
        if input_str == '#':
            continue

        groupss: list[str] = input_str.split(' ')

        group = int(groupss[0])

        if group != -1:
            for group in groupss:
                g_inp = int(group) - 1
                group_id_val = groups[g_inp]['id']
                printChildren()
        else:
            for i in range(0, len(groups)):
                g_inp = i
                group_id_val = groups[i]['id']
                printChildren()

    if choose == '3':
        stat_of_ages(by_program_name=True)

    if choose == '4':
        print('Группы')
        for g in groups:
            i = i + 1
            print(str(i) + ' ' + g['program_name'] + ' ' + g['id'] + " " + g['name'])
        print('Какую группу вывести на печать? ')

        input_str = input()
        if input_str == '#':
            continue

        getListChildrensFromOrderAnyGroups(input_str)

    if choose == '5':
        filename = input("Введи название файла плес")

        if filename == '#':
            continue

        df = pd.read_excel(filename)

        print('Группы')
        for g in groups:
            i = i + 1
            print(str(i) + ' ' + g['program_name'] + ' ' + g['id'] + " " + g['name'])

        input_str = input("Выбери группу ")
        if input_str == '#':
            continue

        g_inp = int(input_str)-1

        print("Статус:", end="")

        for row in df.itertuples():
            if not pandas.isnull(row[2]):
                close_day(row[2].strftime('%Y-%m-%d %H:%M:%S'), row[3], row[4], row[5])
                print("\rСтатус: {0}".format(str(row[2])), end="")

    if choose == '6':
        input_str = input("Группы до какого количества человек Вы хотели бы найти? ")
        if input_str == '#':
            continue
        target_count = int(input_str)
        number_6(target_count)

    if choose == '7':
        find_duplicates()

    if choose == '8':
        stat_of_ages(True)

    #Не надо менять open
    if choose == '9':
        file_exits = os.path.isfile('negative_groups.txt')
        if file_exits:
            f = open('negative_groups.txt', 'r', encoding="utf-8")
            negatve_groups = f.readlines()
            f.close()
            stat_of_ages(by_program_name=True, negative_groups=negatve_groups)
        else:
            print('Файл с шаблоном negative_groups не найден')
            stat_of_ages(by_program_name=True)

    if choose == '10':
        forced_child_adding()
    if choose == '11':
        forced_child_adding(False)

    if choose == '12':
        to_study_from_approve()

    if choose == '13':
        print('Группы для генерации диагностики: \n')
        for g in groups:
            i = i + 1
            print(str(i) + ' ' + g['program_name'] + ' ' + g['id'] + " " + g['name'])

        input_str = input('Выберите группу ')
        if input_str == '#':
            continue

        getDiagnostics(input_str)

    if choose == '14':
        child_search_online()

    if choose == '15':
        print('Группы для генерации диагностики: \n')
        for g in groups:
            i = i + 1
            print(str(i) + ' ' + g['program_name'] + ' ' + g['id'] + " " + g['name'])

        input_str = input('Выберите группу ')
        if input_str == '#':
            continue
        getDiagnostics(input_str, False)



