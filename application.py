from pathlib import Path
import os
import random

import pandas
import pandas as pd
import requests
import json
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt


class App:
    url = "https://booking.dop29.ru/api/user/login"
    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:95.0) Gecko/20100101 Firefox/95.0"
    dir = None
    email = None
    password = None
    YEAR = None
    OUTPUT_DIR = None
    session = None
    use_filter = False
    groups = None
    user = None
    refresh_token = None
    expired_at = None
    access_token = None
    teachers = {}
    selected_groups = None
    max_groups_count = 500
    
    def __init__(self):
        self.dir = os.getcwd()
        file_login = open(self.dir + '\\login.ini', 'r')
        str_login = file_login.read().split('\n')
        self.email = str_login[0]
        self.password = str_login[1]
        self.YEAR = str_login[2]
        self.OUTPUT_DIR = Path(self.dir) / 'output'
        
        if not self.OUTPUT_DIR.exists():
            os.makedirs(self.OUTPUT_DIR)

        self.session = requests.Session()

    def get_save_path(self, filename):
        now = datetime.now()
        str_now = now.strftime('%m-%d-%y')

        folder_path = self.OUTPUT_DIR / str_now

        if not folder_path.exists():
            os.makedirs(folder_path)

        return folder_path / filename

    def auth(self):
        try:
            r = self.session.post(self.url, headers={
                'Host': 'booking.dop29.ru',
                'User-Agent': self.user_agent,
                'Accept': '*/*',
                'Accept-Language': 'ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3',
                'Accept-Encoding': 'gzip, deflate, br',
                'Content-Type': 'application/json',
                'X-Requested-With': 'XMLHttpRequest',
                'Content-Length': '63',
                'Origin': 'https://booking.dop29.ru',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Referer': 'https://booking.dop29.ru/admin/',
                'Sec-Fetch-Dest': 'empty',
                'Sec-Fetch-Mode': 'cors',
                'Sec-Fetch-Site': 'same-origin',
                'TE': 'trailers',
            }, data='{"email": "' + self.email + '", "password": "' + self.password + '"}')
            self.session.headers.update({'Referer': 'https://booking.dop29.ru/admin/'})
            self.session.headers.update({'User-Agent': self.user_agent})

            text_buf = r.text
            json_string = json.loads(text_buf)

            if json_string['err_code'] == 400:
                return json_string['errors'][0]['msg']
            else:
                self.access_token = json_string['data']['access_token']
                self.expired_at = json_string['data']['expired_at']
                self.refresh_token = json_string['data']['refresh_token']

                self.user = json_string['data']['user']

                return 0
        except Exception as e:
            return e

    def get_all_groups(self):
        self.headers = {
            'Host': 'booking.dop29.ru',
            'User-Agent': self.user_agent,
            'Accept': '*/*',
            'Accept-Language': 'ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3',
            'Accept-Encoding': 'gzip, deflate, br',
            'Authorization': 'Bearer ' + self.access_token,
            'X-REQUEST-ID': '7bd411c3-54ce-4bba-9ee1-7c5091da6d1a',
            'X-Requested-With': 'XMLHttpRequest',
            'DNT': '1',
            'Connection': 'keep-alive',
            'Referer': 'https://booking.dop29.ru/admin/',
            'Cookie': 'io=lVluIaMvSTa4ImFmB5C9',
            'Sec-Fetch-Dest': 'empty',
            'Sec-Fetch-Mode': 'cors',
            'Sec-Fetch-Site': 'same-origin',
            'TE': 'trailers'
        }

        new_url = 'https://booking.dop29.ru/api/rest/eventGroups?_dc=1641896017213&page=1&start=0&length=25&extFilters=[{"property":"is_deleted","value":"0","comparison":"eq"},{"property":"event.is_deleted","value":"N","comparison":"eq"}]&format=attendance&length=' + str(
            self.max_groups_count)
        r = self.session.get(new_url, headers=self.headers)

        b = json.loads(r.text)
        self.groups = b['data']

        if int(b['recordsFiltered']) > len(self.groups):

            new_url = 'https://booking.dop29.ru/api/rest/eventGroups?_dc=1641896017213&page=1&start=0&length=25&extFilters=[{"property":"is_deleted","value":"0","comparison":"eq"},{"property":"event.is_deleted","value":"N","comparison":"eq"}]&format=attendance&length=' + str(
                self.max_groups_count) + '&page=2&start=' + str(len(self.groups))
            r = self.session.get(new_url, headers=self.headers)
            b = json.loads(r.text)
            self.groups.extend(b['data'])

        return "Загружено {0} из {1}".format(len(self.groups), int(b['recordsFiltered']))

    def get_teachers(self):
        self.teachers = {}
        groups_id = []
        with open(Path(self.dir) / 'groups.ini', 'r', encoding='utf-8') as f:
            groups_id = f.read().split('\n')

        if self.use_filter:
            for g in self.groups:
                if g['id'] not in groups_id:
                    self.groups.remove(g)

        for g in self.groups:
            if self.use_filter:
                if g['id'] in groups_id:

                    self.fill_teachers(g)
                else:
                    continue
            else:
                self.fill_teachers(g)

        keys = []
        for t in self.teachers:
            keys.append(t)

        return keys

    def fill_teachers(self, g):
        if self.teachers.get(g["teacher"]) is not None:
            self.teachers[g["teacher"]].append(g)
        else:
            self.teachers[g["teacher"]] = []
            self.teachers[g["teacher"]].append(g)

    def select_groups(self, teacher_groups):
        groups = []
        for t in teacher_groups:
            groups.extend(self.teachers[t])

        self.groups = groups

    def printChildren(self):
        ids = self.selected_groups.strip().split(' ')
        ids = [int(i) for i in ids]
        ret_msg = []
        for i in ids:
            group_id_val = self.groups[i]['id']
            print('Выбрана группа ' + self.groups[i]['program_name'] + ' ' + self.groups[i]['name'])
            list_childrens = self.get_childrens(group_id_val=group_id_val)
            f = open(self.get_save_path(f"Список группы {self.groups[i]['program_name']} {self.groups[i]['name']}.txt"), 'w',
                     encoding="utf-8")
            for c in list_childrens:
                line = (f'{c['kid_last_name']} {c['kid_first_name']} {c['kid_patro_name']}\t'
                        f'{c['kid_birthday'].replace('-', '.')}\t{c['kid_age']}\n')
                f.write(line)
            f.close()
            ret_msg.append(f"Созданы: {self.get_save_path(f"Список группы {self.groups[i]['program_name']} "
                                                          f"{self.groups[i]['name']}.txt")}\n")
        return ret_msg

    def get_childrens(self, group_id_val):
        new_url = 'https://booking.dop29.ru/api/attendance/members/get?_dc=1641896197594&page=1&start=0&length=25&extFilters=[{"property":"group_id","value":"' + str(
            group_id_val) + '"},{"property":"academic_year_id","value":"' + str(
            self.YEAR) + '"},{"property":"dateStart","value":"' + self.YEAR + '-12-01 00:00:00"},{"property":"dateEnd","value":"' + self.YEAR + '-12-31 23:59:59"}]'
        r = self.session.get(new_url, headers=self.headers)
        b = json.loads(r.text)
        list_childrens = b['data']
        new_list_childrens = []
        for i in range(0, len(list_childrens)):
            if list_childrens[i]['type_active'] == 1:
                new_list_childrens.append(list_childrens[i])
        return new_list_childrens

    def stat_of_ages(self, unique=False, confirmed=False, by_program_name=False, negative_groups=[], status_wiget = None):
        ages = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0, 8: 0, 9: 0, 10: 0, 11: 0, 12: 0, 13: 0, 14: 0, 15: 0,
                16: 0, 17: 0, 18: 0, 19: 0, 20: 0, 21: 0}
        stupid_girls_by_ages = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0, 8: 0, 9: 0, 10: 0, 11: 0, 12: 0, 13: 0,
                                14: 0, 15: 0, 16: 0, 17: 0, 18: 0, 19: 0, 20: 0, 21: 0}
        sum_girls = 0
        sum_childs = 0
        error_sex = 0
        error_child = []

        unique_childs = set()
        repeated_childs = 0
        repeteds = 0
        repeated_childs_id = []

        ages_of_sections = {}  # секция: [мальчики, девочки]
        for i in range(0, len(self.groups)):

            if len(negative_groups) > 0:
                negative_check = [ind for ind in negative_groups if ind in self.groups[i]['name'].lower()]
                if len(negative_check) != 0:
                    print('\r Проигнорированна группа ' + self.groups[i]['name'] + '\n')
                    continue

            g_inp = i
            group_id_val = self.groups[i]['id']
            event_id = self.groups[i]['event_id']


            list_childrens = self.get_childrens(group_id_val=group_id_val)
            if by_program_name:
                section = self.groups[i]['program_name']
            else:
                section = self.get_section(event_id)

            if section not in ages_of_sections:
                ages_of_sections[section] = [0], [0], {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0, 8: 0, 9: 0,
                                                       10: 0, 11: 0, 12: 0, 13: 0, 14: 0, 15: 0, 16: 0, 17: 0, 18: 0,
                                                       19: 0, 20: 0, 21: 0}  # Мальчики, девочки

            iterator_childrens = 0

            for c in list_childrens:
                if unique:
                    if c['kid_id'] in unique_childs:
                        repeteds += 1
                        if c['kid_id'] not in repeated_childs_id:
                            repeated_childs += 1
                            repeated_childs_id.append(c['kid_id'])
                        continue
                    else:
                        unique_childs.add(c['kid_id'])

                ages[c['kid_age']] += 1
                ages_of_sections[section][2][c['kid_age']] += 1
                c_info = self.get_all_info_child(c['kid_id'])
                sum_childs += 1

                try:
                    if c_info['sex'] == 'W':
                        sum_girls += 1
                        ages_of_sections[section][1][0] += 1
                        stupid_girls_by_ages[c['kid_age']] += 1
                    else:
                        ages_of_sections[section][0][0] += 1
                except:
                    error_sex += 1
                    error_child.append(
                        "{0} {1} {2} {3}".format(c['kid_last_name'], c['kid_first_name'], c['kid_patro_name'],
                                                 c['kid_birthday']))

                iterator_childrens += 1

                if status_wiget is not None:
                    status_wiget.values.add(f"Выбрана группа {self.groups[g_inp]['program_name']} {self.groups[g_inp]['name']}")

        if not unique:
            f = open(self.get_save_path("Статистика по возрастам.txt"), "w", encoding='utf-8')
        else:
            f = open(self.get_save_path("Статистика по возрастам УНИКАЛЬНЫЕ.txt"), "w", encoding='utf-8')
        for i in range(0, 19):
            if ages[i] == 0:
                continue
            else:
                f.write(str(i) + " лет " + str(ages[i]) + f" человек; {stupid_girls_by_ages[i]} из них девочек \n")

        f.write("Всего: {0}, из них девочек: {1}".format(sum_childs, sum_girls))
        f.write("\nНе удалось получить информацию у {0} человек:".format(error_sex))
        for l in error_child:
            f.write(l)
        for key, value in ages_of_sections.items():
            f.write("\n\nНаправленность: {0}, М {1}, Ж {2}\n\n".format(key, value[0][0], value[1][0]))

            for i in range(0, 19):
                if value[2][i] == 0:
                    continue
                else:
                    f.write(str(i) + " лет " + str(value[2][i]) + " человек\n")

        if unique:
            f.write("\nПовторов всего: " + str(repeteds))
            f.write("\nПовторов детей: " + str(repeated_childs))
            f.write('\n' + str(repeated_childs_id))

        f.close()
        if not unique:
            return [f"Файл {self.get_save_path("Статистика по возрастам.txt")} создан успешно!"]
        else:
            return [f"Файл {self.get_save_path("Статистика по возрастам УНИКАЛЬНЫЕ.txt")} создан успешно!"]

    def get_section(self, event_id):
        new_url = 'https://booking.dop29.ru/api/rest/events/{0}?_dc=1705994318220'.format(event_id)
        r = self.session.get(new_url, headers=self.headers)
        b = json.loads(r.text)
        try:
            return b['data'][0]['section']
        except:
            return []

    def get_all_info_child(self, id):
        new_url = "https://booking.dop29.ru/api/rest/kid/{0}?_dc=1704971612231".format(id)
        r = self.session.get(new_url, headers=self.headers)
        b = json.loads(r.text)
        try:
            child = b['data'][0]
            return child
        except:
            return []

    def getListChildrensFromOrderAnyGroups(self):
        ret_str = []
        ids = self.selected_groups.strip().split(' ')
        ids = [int(i) for i in ids]

        for i in ids:
            ret_str.append(self.getFileListChildrensFromOrder(self.groups[i]))

        return ret_str

    def getFileListChildrensFromOrder(self, group):
        group_id_val = group['id']

        new_url = 'https://booking.dop29.ru/api/rest/order?_dc=1695285515100&page=1&start=0&length=25&extFilters=[{"property":"fact_academic_year_id","value":' + self.YEAR + ',"comparison":"eq"},{"property":"event_id","value":' + \
                  group['event_id'] + ',"comparison":"eq"},{"property":"fact_group_id","value":"' + str(
            group_id_val) + '","comparison":"eq"},{"property":"state","value":["approve"],"comparison":"in"}]'

        r = self.session.get(new_url, headers=self.headers)
        b = json.loads(r.text)
        list_childrens = b['data']
        list_names = []
        for children in list_childrens:
            url_child = 'https://booking.dop29.ru/api/rest/kid/' + children['kid_id']
            r = self.session.get(url_child, headers=self.headers)
            child = json.loads(r.text)['data'][0]
            list_names.append(child['last_name'] + " " + child['first_name'] + " " + child['patro_name'])

        filename = self.get_save_path(f"Подтверждённые заявки {group['program_name']} {group['name']}.txt")

        file = open(filename, 'w', encoding="utf-8")
        file.write("\n".join(list_names))
        file.close()

        return f"Файл {filename} создан!"

    def up_close_day(self, filename, group):
        df = pd.read_excel(filename)

        for row in df.itertuples():
            if not pandas.isnull(row[2]):
                self.close_day(row[2].strftime('%Y-%m-%d'), row[3], row[4], row[5], group)
        return ['Внесение информации завершено. Проверьте в навигаторе дни!']

    # Типы занятий
    # Практическая работа 9732
    # Учебное 7198
    # Дистанционное 3022
    #
    def close_day(self, date, theme, type, description, group):
        percentagevisits = 80
        group_id_val = group['id']
        group_id = group['id']

        childrens = self.get_childrens(group_id_val)

        for c in childrens:
            visit = random.randint(0, 100)
            if visit < percentagevisits:
                visit = True
            else:
                visit = False
            kid_id = c['kid_id']
            day_url_POST = "https://booking.dop29.ru/api/attendance/save"  # date group_id kid_id value:true
            payload = {'date': date, 'group_id': group_id, 'kid_id': kid_id, 'value': visit}
            r = self.session.post(url=day_url_POST, headers=self.headers, data=json.dumps(payload))

        KTP_url_POST = "https://booking.dop29.ru/api/event-group-lessons/upsert"  # date group_id types:list description

        payload = json.loads(json.dumps({'data': {'date': date + ' 00:00:00', 'group_id': group_id, 'theme': theme,
                                                  'types': [str(int(type))], 'description': description}}))
        r = self.session.post(url=KTP_url_POST, headers=self.headers, json=payload)

    def number_6(self, target_sum):
        problem_groups = []
        for g in self.groups:
            id = g['id']
            group_id_val = id
            childrens = self.get_childrens(group_id_val)
            if len(childrens) < target_sum and len(childrens) != 0:
                problem_groups.append(
                    "Группа: {0}, {1} человек!".format(g['program_name'] + " " + g['name'], len(childrens)))

        f = open(self.get_save_path("ПРОБЛЕМНЫЕ ГРУППЫ.txt"), "w")
        for g in problem_groups:
            f.write(g + '\n')
        f.close()
        return [f'Файл {self.get_save_path("ПРОБЛЕМНЫЕ ГРУППЫ.txt")} создан!']

    def find_duplicates(self):
        childs_and_groups = {}
        for g in self.groups:
            id = g['id']
            group_id_val = id
            childs = self.get_childrens(group_id_val)

            for child in childs:
                full_name = child['kid_last_name'] + ' ' + child['kid_first_name'] + ' ' + child['kid_patro_name']
                group_name = g['program_name'] + ' ' + g['name']

                if full_name in childs_and_groups:
                    childs_and_groups[full_name].append(group_name)
                else:
                    childs_and_groups[full_name] = []
                    childs_and_groups[full_name].append(group_name)

                pass

        duplicated = {}
        for key in childs_and_groups:
            if len(childs_and_groups[key]) > 1:
                duplicated[key] = childs_and_groups[key]

        f = open(self.get_save_path("Дубликаты.txt"), "w")
        for key in duplicated:

            str_groups = '\n'
            for i in range(len(duplicated[key])):
                str_groups += '\t' + duplicated[key][i] + '\n'
            str_groups += '\n'

            f.write('Групп ' + str(len(duplicated[key])) + ' ' + key + ' ' + str_groups + '\n')

        f.write('\n\n Всего детей: ' + str(len(duplicated)))
        f.close()
        return[f'Файл {self.get_save_path("Дубликаты.txt")} создан', f'Всего детей: {len(duplicated)}']

    def count_child_by_program(self, negative_groups_filename=None):
        if negative_groups_filename:
            f = open(negative_groups_filename, 'r', encoding="utf-8")
            negatve_groups = f.readlines()
            f.close()
            return self.stat_of_ages(by_program_name=True, negative_groups=negatve_groups)

        else:
            return self.stat_of_ages(by_program_name=True)

    def forced_child_adding(self, filename, in_group=True):
        # Файл с детьми которых нужно зачислить в группы
        # Текстовый
        # ФИО таб часть program_name

        # Файл с детьми которых нужно зачислить в мероприятия
        # Текстовый
        # ИД мероприяти таб дата-время
        # ФИО таб описание

        all_childrens = []

        if in_group:
            for g in self.groups:
                childs = self.get_childrens_by_group_id(g['id'])
                all_childrens.extend(childs)
                also_childs = self.getListChildrensFromOrder(g)
                all_childrens.extend(also_childs)

        f = open(filename, 'r', encoding='utf-8')
        rows = f.readlines()
        f.close()

        activity_id = 0
        date = ""
        if not in_group:
            parts = rows[0].split('\t')
            activity_id = int(parts[0])
            date = parts[1]
            del (rows[0])

        prerared_info = []
        for r in rows:
            parts = r.split('\t')
            if parts[0] != '':
                prerared_info.append(parts)

        actual_kids_for_adding = []
        self.search_and_processing_child(activity_id, actual_kids_for_adding, all_childrens, date, in_group,
                                         prerared_info)

    def search_and_processing_child(self, activity_id, actual_kids_for_adding, all_childrens, date, in_group,
                                    prerared_info):
        for info in prerared_info:
            full_name = info[0]
            try:
                if info[1] is not None:
                    adding_info = info[1]
                else:
                    adding_info = ''
            except IndexError:
                adding_info = ''

            print('Поиск {0}'.format(full_name))
            target_url = f'https://booking.dop29.ru/api/rest/safe/kid?_dc=1714046462894&special=1&page=1&start=0&length=20&extFilters=[{{"property":"fio","value":"{full_name}","comparison":"manual","type":null}}]'
            r = self.session.get(url=target_url, headers=self.headers)
            b = json.loads(r.text)

            if b['err_code'] != 0 or len(b['data']) == 0:
                print('Не найдено!')
                continue
            os.system('cls')
            print('Найдены следующие дети:')

            for i in range(len(b['data'])):
                print(f"{i} {b['data'][i]['fio']} {b['data'][i]['birthday']} {b['data'][i]['approve_org_caption']}")

            if not in_group:
                print(f'Описание: {adding_info}')

            choose = int(input('Выберите индекс ребёнка для добавления или напишите -1 для пропуска'))
            if choose == -1:
                continue

            actual_kids_for_adding.append(b['data'][choose])

            self.processing_child(activity_id, adding_info, all_childrens, b, choose, date, in_group)

    def processing_child(self, activity_id, adding_info, all_childrens, b, choose, date, in_group):
        if in_group:
            print("Веберите идекс группы для зачисления")
            for i in range(len(self.groups)):
                if adding_info.lower().rstrip() in self.groups[i]['program_name'].lower().rstrip():
                    print(f"{i} {self.groups[i]['program_name']} {self.groups[i]['id']} {self.groups[i]['name']}")

            if b['data'][choose]['id'] in [kid['kid_id'] for kid in all_childrens]:
                print('УЖЕ ДОБАВЛЕН В КАКОЙ-ТО ГРУППЕ')

            group_index = int(input('Выбранная группа: '))
            if group_index != -1:
                print(
                    f"{self.groups[group_index]['program_name']} {self.groups[group_index]['id']} {self.groups[group_index]['name']}")
                self.adding_order(b['data'][choose], self.groups[group_index])
            else:
                print('Пропущен!')
        else:
            self.adding_activity_order(b['data'][choose], activity_id, date)
        pass

    def get_childrens_by_group_id(self, group_id):
        new_url = 'https://booking.dop29.ru/api/attendance/members/get?_dc=1641896197594&page=1&start=0&length=25&extFilters=[{"property":"group_id","value":"' + str(
            group_id) + '"},{"property":"academic_year_id","value":"' + str(
            self.YEAR) + '"},{"property":"dateStart","value":"' + self.YEAR + '-12-01 00:00:00"},{"property":"dateEnd","value":"' + self.YEAR + '-12-31 23:59:59"}]'
        r = self.session.get(new_url, headers=self.headers)
        b = json.loads(r.text)
        list_childrens = b['data']
        new_list_childrens = []
        for i in range(0, len(list_childrens)):
            if list_childrens[i]['type_active'] == 1:
                new_list_childrens.append(list_childrens[i])
        return new_list_childrens

    def getListChildrensFromOrder(self, group):

        group_id_val = group['id']

        new_url = 'https://booking.dop29.ru/api/rest/order?_dc=1695285515100&page=1&start=0&length=25&extFilters=[{"property":"fact_academic_year_id","value":' + self.YEAR + ',"comparison":"eq"},{"property":"event_id","value":' + \
                  group['event_id'] + ',"comparison":"eq"},{"property":"fact_group_id","value":"' + str(
            group_id_val) + '","comparison":"eq"},{"property":"state","value":["approve"],"comparison":"in"}]'

        r = self.session.get(new_url, headers=self.headers)
        b = json.loads(r.text)
        list_childrens = b['data']
        return list_childrens

    def adding_order(self, child, group):
        json_string = {"data":
                           {"event_id": group['event_id'],
                            "state": "initial",
                            "certificate_number": "нет",
                            "decree_enrollment_number": "нет",
                            "decree_deduction_number": "нет",
                            "program_is_pfdod": False,
                            "kid_is_approved": False,
                            "is_online_payments_allowed": False,
                            "academic_year_id": self.YEAR,
                            "certificate_certificate_number": "",
                            "rpgu_deadline_date": None,
                            "kid_birthday": None,
                            "deadline": None,
                            "created_ts": None,
                            "date_enroll": None,
                            "date_deduct": None,
                            "rpgu_overdue_deadline": False,
                            "group_id": f"{group['id']}",
                            "kid_id": f"{child['id']}",
                            "site_user_id": f"{child['site_user_id']}"
                            }}
        new_url = "https://booking.dop29.ru/api/rest/order?_dc=1714059370156"
        payload = json.loads(json.dumps(json_string))
        r = self.session.post(url=new_url, headers=self.headers, json=payload)
        b = json.loads(r.text)

        if b['err_code'] == 0:
            return('УСПЕХ!')
        else:
            return(b['errors'][0]['msg'])

    def adding_activity_order(self, child, activity_id, date, state='approve'):
        json_string = {"data":
            {
                "activity_id": activity_id,
                "date": date,
                "site_user_id": f"{child['site_user_id']}",
                "kid_id": f"{child['id']}",
                "state": state
            }}
        new_url = "https://booking.dop29.ru/api/rest/activityOrder?_dc=1714066258891"
        payload = json.loads(json.dumps(json_string))
        r = self.session.post(url=new_url, headers=self.headers, json=payload)
        b = json.loads(r.text)
        if b['err_code'] == 0:
            return('УСПЕХ!')
        else:
            return(b['errors'][0]['msg'])

    def to_study_from_approve(self):
        target_url = f'https://booking.dop29.ru/api/rest/order?_dc=1714828457975&page=1&start=0&length={self.max_groups_count}&extFilters=[{{"property":"fact_academic_year_id","value":{self.YEAR},"comparison":"eq"}}]'
        r = self.session.get(url=target_url, headers=self.headers)
        b = json.loads(r.text)
        approving = [a for a in b['data'] if a['state_grid'] == 'approve']
        print(f'Найдено {len(approving)} подтверждённых заявок для подтверждения обучения')
        date_signing = ""  # Дата приказа
        date_start = ""  # Начало обучения
        decree_number = ""  # Номер приказа
        financing_source = "1"  # Бюджет 1
        id = ""  # id заявки ребёнка

        print('=' * 20)
        date_signing = input('Дата приказа в формате ГГ-ММ-ДД! ')
        date_start = input('Дата начала обучения в формате ГГ-ММ-ДД!')
        decree_number = input('Номер приказа ')

        for a in approving:
            g = next((x for x in self.groups if x['id'] == a['group_id']), None)
            if g is None:
                continue

            print(f"Ребёнок {a['kid_last_name']} {a['kid_first_name']} {g['program_name']} {g['name']}")
            choose = input('1 - Принять; 0 - пропустить;\n')
            if choose == '1':
                id = a['id']

                target_url = f'https://booking.dop29.ru/api/studyRequest'
                json_string = {"data":
                    {
                        "comment": "",
                        "date_signing": date_signing,
                        "date_start": date_start,
                        "decree_number": decree_number,
                        "financing_source": "1",
                        "id": id
                    }}
                payload = json.loads(json.dumps(json_string))
                r = self.session.post(url=target_url, headers=self.headers, json=payload)
                b = json.loads(r.text)
                if b['err_code'] == 0:
                    print('УСПЕХ!')
                else:
                    b['errors'][0]['msg']
            else:
                continue

    def getDiagnostics(self, existing=True):
        groups = self.selected_groups
        diagnostics_sums = {}
        if ' ' in groups:
            groups = groups.split(' ')
            groups.remove('')
            for group in groups:
                self.generateDiagnostic(int(group) - 1,diagnostics_sums, existing=existing)
        else:
            self.generateDiagnostic(int(groups) - 1, diagnostics_sums, existing=existing)
        f = open(self.get_save_path('Диагностика суммы.txt'), 'w', encoding='utf-8')
        for key, value in diagnostics_sums.items():
            f.write(f"{key} Высокий: {value['high']} Средний: {value['middle']}\n")
        f.close()
        return [f'Файлы созданы']

    def generateDiagnostic(self, group, diagnostics_sums, existing=True):

        group_id_val = self.groups[int(group)]['id']
        childrens = self.get_childrens(group_id_val)
        # Высокий уровень 30%, средний 70% для выходной
        list_fio = [f"{c['kid_last_name']} {c['kid_first_name']} {c['kid_patro_name']}" for c in childrens]

        table = []
        summary = self.generate_data(diagnostics_sums, group, self.groups, list_fio, table, existing=existing)

        header_table = [['#п/п', 'ФИО обучающегося', 'уровень знаний', '', ''],
                        ['', '', 'низкий', 'средний', 'высокий']]

        header_table.extend(table)
        table = header_table

        doc = self.create_document()

        if existing:
            parts = [(f"Выходная диагностика {self.groups[int(group)]['program_name']} {self.groups[int(group)]['name']}", True)]
        else:
            parts = [(f"Входная диагностика {self.groups[int(group)]['program_name']} {self.groups[int(group)]['name']}", True)]
        self.add_paragraph(doc, parts, font_size=14, alignment='center')

        t = self.add_table(doc, table)
        self.merge_cells_with_content(t, 0, 0, 1, 0)
        self.merge_cells_with_content(t, 0, 1, 1, 1)
        self.merge_cells_with_content(t, 0, 2, 0, 4)

        parts = [(f"Итого: {summary[0]} низкий, {summary[1]} средний, {summary[2]} высокий", False)]
        self.add_paragraph(doc, parts, font_size=14, alignment='center')

        self.add_page_break(doc)

        parts = [(f"Критерии оценки", True)]
        self.add_paragraph(doc, parts, font_size=14, alignment='center')

        table = [['Уровень знаний и умений', 'Низкий уровень', 'Средний уровень', 'Высокий уровень'],
                 ['Теоретические знания', 'Определяются по результатам собеседования', '', ''],
                 ['Практические умения и навыки',
                  'Ребенок не смог выполнить задание без помощи педагога или работал самостоятельно, но задание выполнено не верно',
                  'Задание выполнено хорошо, но ребенок задавал вопросы в процессе выполнения',
                  'Задание выполнено самостоятельно, быстро и качественно'],
                 ['Личностные качества',
                  'Определяются в результате педагогического наблюдения в процессе выполнения задания', '', ''],
                 ]

        t = self.add_table(doc, table)
        self.merge_cells_with_content(t, 1, 1, 1, 3)
        self.merge_cells_with_content(t, 3, 1, 3, 3)

        parts = [(f"Форма определения уровня освоения программы:", True),
                 (" педагогическое наблюдение, собеседование, анализ практической работы, результат проекта.", False)]
        self.add_paragraph(doc, parts, font_size=14, alignment='justify')

        if existing:
            self.save_document(doc,
                          self.get_save_path(
                              f"Выходная диагностика {self.groups[int(group)]['program_name']} {self.groups[int(group)]['name']}.docx"))
            return [f'Файл {self.get_save_path(
                              f"Выходная диагностика {self.groups[int(group)]['program_name']} {self.groups[int(group)]['name']}.docx")} создан']
        else:
            self.save_document(doc,
                          self.get_save_path(
                              f"Входная диагностика {self.groups[int(group)]['program_name']} {self.groups[int(group)]['name']}.docx"))
            return [f'Файл {self.get_save_path(
                              f"Входная диагностика {self.groups[int(group)]['program_name']} {self.groups[int(group)]['name']}.docx")} создан']

    def create_document(self):
        """Создаёт новый документ Word."""
        return Document()

    def add_paragraph(self, doc, parts, font_size=14, alignment=None):
        """
        Добавляет абзац с текстом в документ, поддерживая отдельные жирные слова.
        :param doc: документ
        :param parts: список кортежей (текст, жирный), например, [("Hello", False), ("World", True)]
        :param font_size: размер шрифта
        :param alignment: выравнивание текста (None, 'center', 'left', 'right', 'justify')
        """
        paragraph = doc.add_paragraph()
        for text, bold in parts:
            run = paragraph.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)

        if alignment:
            if alignment == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'left':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif alignment == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == 'justify':
                # Выравнивание по обеим сторонам (приближение к выравниванию по ширине)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return paragraph

    def set_cell_border(self, cell, **kwargs):
        """
        Устанавливает границы для ячейки.
        :param cell: ячейка таблицы
        :param kwargs: параметры границ
        """
        tcPr = cell._element.get_or_add_tcPr()

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = tcPr.find(qn(f'w:{border_name}'))
            if border is None:
                border = OxmlElement(f'w:{border_name}')
                tcPr.append(border)
            for attr, value in kwargs.items():
                border.set(qn(f'w:{attr}'), str(value))

    def add_table(self, doc, data):
        """
        Добавляет таблицу в документ с границами.
        :param doc: документ
        :param data: список списков, представляющий строки и столбцы таблицы
        """
        table = doc.add_table(rows=len(data), cols=len(data[0]))

        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table_cell = table.cell(i, j)
                table_cell.text = str(cell)
                self.set_cell_border(table_cell, val="single", sz="4", space="0", color="000000")

        # Автоматически подгоняем ширину столбцов под текст
        for col in table.columns:
            max_length = max(len(cell.text) for cell in col.cells)
            for cell in col.cells:
                cell.width = Inches(0.15 * max_length)  # Установка ширины в зависимости от длины текста
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return table

    def merge_cells_with_content(self, table, start_row, start_col, end_row, end_col):
        """
        Объединяет ячейки таблицы от start_row, start_col до end_row, end_col, если одна из ячеек заполнена, а другие пустые.
        :param table: таблица
        :param start_row: начальная строка
        :param start_col: начальный столбец
        :param end_row: конечная строка
        :param end_col: конечный столбец
        """
        # Проверка содержимого ячеек
        content = None
        for row in range(start_row, end_row + 1):
            for col in range(start_col, end_col + 1):
                cell_text = table.cell(row, col).text.strip()
                if cell_text:
                    if content and content != cell_text:
                        raise ValueError("Конфликтующее содержимое ячеек")
                    content = cell_text

        # Объединение ячеек
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        start_cell.merge(end_cell)

        if content:
            start_cell.text = content

    def center_text(self, cell):
        """
        Выравнивает текст в ячейке по центру.
        :param cell: ячейка
        """
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(14)  # Установка размера шрифта

    def add_line_break(self, doc):
        """Добавляет разрыв строки в документ."""
        doc.add_paragraph().add_run().add_break()

    def add_page_break(self, doc):
        """Добавляет разрыв страницы в документ."""
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save_document(self, doc, filename):
        """Сохраняет документ под указанным именем файла."""
        doc.save(filename)

    def generate_data(self, diagnostics_sums, group, groups, list_fio, table, existing=True):
        summary = [0, 0, 0]
        for i in range(len(list_fio)):
            r = random.randint(0, 100)
            if r >= 70:
                if existing:
                    table.append([i + 1, list_fio[i], "", "", "+"])
                else:
                    table.append([i + 1, list_fio[i], "", "+", ""])
                if f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}" not in \
                        diagnostics_sums.keys():
                    diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"] = \
                        {'high': 0, 'middle': 0}
                diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"]['high'] += 1

                summary[2] += 1
            else:
                if existing:
                    table.append([i + 1, list_fio[i], "", "+", ""])
                else:
                    table.append([i + 1, list_fio[i], "+", "", ""])

                if f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}" not in \
                        diagnostics_sums.keys():
                    diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"] = \
                        {'high': 0, 'middle': 0}
                diagnostics_sums[f"{groups[int(group)]['teacher']} {groups[int(group)]['program_name']}"]['middle'] += 1

                summary[1] += 1
        return summary

    def child_search_online(self, full_name):
        b = self.search_child_by_fio(full_name)

        ret_vals = []

        ret_vals.append('Найдены следующие дети:')

        if b == 'Не найдено!':
            return ['Не найдено']

        for i in range(len(b['data'])):
            ret_vals.append(f"{i} {b['data'][i]['fio']} {b['data'][i]['birthday']} {b['data'][i]['approve_org_caption']}")

        return ret_vals

    def search_child_by_fio(self, full_name):
        target_url = f'https://booking.dop29.ru/api/rest/safe/kid?_dc=1714046462894&special=1&page=1&start=0&length=20&extFilters=[{{"property":"fio","value":"{full_name}","comparison":"manual","type":null}}]'
        r = self.session.get(url=target_url, headers=self.headers)
        b = json.loads(r.text)
        if b['err_code'] != 0 or len(b['data']) == 0:
            return 'Не найдено!'
        return b
